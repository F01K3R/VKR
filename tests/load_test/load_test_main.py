import time
import psutil
import os
import sys
import threading
import random
import string
from multiprocessing import Pool
from docx import Document
from docx.shared import Pt
from main import process_file  # Импортируем process_file из main.py


def generate_random_text(length=500):
    """Генерирует случайный текст заданной длины."""
    return ''.join(random.choices(string.ascii_letters + string.digits + " ", k=length))


def generate_test_doc(args):
    """Генерирует тестовый документ с уникальным содержимым."""
    filename, num_paragraphs, include_references = args
    if os.path.exists(filename):
        return filename
    doc = Document()
    doc.add_heading("Оглавление", level=1)
    doc.add_heading("Введение", level=1)
    for i in range(num_paragraphs):
        p = doc.add_paragraph(f"Тестовый параграф {i}: {generate_random_text()}")
        p.style.font.name = "Times New Roman"
        p.style.font.size = Pt(14)
    doc.add_heading("Заключение", level=1)
    if include_references:
        doc.add_paragraph("Список литературы")
        doc.add_paragraph(
            f"Иванов И.И. Статья // Журнал. {random.randint(2010, 2023)}. № {random.randint(1, 10)}. С. 1-5.")
    doc.save(filename)
    return filename


def monitor_resources(stop_event, interval=0.5):
    """Мониторит использование ресурсов в реальном времени."""
    process = psutil.Process()
    print("\nМониторинг ресурсов в реальном времени (Ctrl+C или кнопка Stop в PyCharm для остановки):")
    print("Time (s) | CPU (%) | Memory (%)")
    start_time = time.time()
    while not stop_event.is_set():
        cpu_usage = psutil.cpu_percent(interval=None)
        memory_usage = process.memory_percent()
        elapsed_time = time.time() - start_time
        print(f"{elapsed_time:.2f}s    | {cpu_usage:6.1f} | {memory_usage:6.2f}")
        time.sleep(interval)


def infinite_load_test(num_files=10, base_paragraphs=100, num_processes=4):
    """Бесконечное тестирование с N различными файлами."""
    test_files_dir = "test_files"
    os.makedirs(test_files_dir, exist_ok=True)

    # Подготовка аргументов для параллельной генерации файлов
    file_args = []
    print(f"Проверка и генерация {num_files} тестовых файлов параллельно с {num_processes} процессами...")
    start_time = time.time()
    for i in range(num_files):
        num_paragraphs = base_paragraphs + random.randint(-50, 50)
        file_path = os.path.join(test_files_dir, f"test_file_{i}.docx")
        file_args.append((file_path, num_paragraphs, True))

    with Pool(processes=num_processes) as pool:
        file_paths = pool.map(generate_test_doc, file_args)

    gen_time = time.time() - start_time
    print(f"Генерация/проверка завершена за {gen_time:.2f} секунд")

    stop_event = threading.Event()
    monitor_thread = threading.Thread(target=monitor_resources, args=(stop_event,))
    monitor_thread.start()

    iteration = 0
    pool = None
    try:
        while not stop_event.is_set():
            iteration += 1
            print(f"\nИтерация {iteration}: Тестирование {num_files} файлов")
            start_time = time.time()
            pool = Pool(processes=8)
            results_list = pool.map(process_file, file_paths)  # Используем process_file из main.py
            pool.close()
            pool.join()
            pool = None
            total_time = time.time() - start_time

            # Собираем времена обработки всех файлов
            file_times = [result["time"] for result in results_list]
            avg_time_per_file = sum(file_times) / len(file_times) if file_times else 0

            # Вывод результатов
            for result in results_list:
                file_path = result["file_path"]
                file_results = result["results"]
                processing_time = result["time"]
                print(f"File {file_path} (время обработки: {processing_time:.2f}s):")
                for check, res in file_results.items():
                    print(f"  {check}: {res}")
            print(f"Total Time for {num_files} files: {total_time:.2f}s")
            print(f"Average Time per File: {avg_time_per_file:.2f}s")

    except KeyboardInterrupt:
        print("\nОстановка тестирования через KeyboardInterrupt...")
        stop_event.set()

    finally:
        if pool:
            pool.terminate()
            pool.join()
        stop_event.set()
        monitor_thread.join()
        print("Тестирование завершено.")


if __name__ == "__main__":
    infinite_load_test(num_files=10000, base_paragraphs=100, num_processes=8)