import pytest
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from modules.formatting import FormattingCheck  # Ваш класс с проверками

# Вспомогательная функция для создания документа с форматированием
def create_test_document(font="Times New Roman", font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=0):
    doc = Document()
    para = doc.add_paragraph("Test text")
    run = para.add_run("Sample run")
    run.font.name = font
    run.font.size = Pt(font_size)
    para.paragraph_format.alignment = alignment
    para.paragraph_format.first_line_indent = Inches(first_line_indent)  # Добавляем отступ первой строки
    return doc


# Тесты

@pytest.fixture
def formatting_check():
    return FormattingCheck()

# Тест 1: Проверка использования неправильного шрифта
def test_wrong_font(formatting_check):
    doc = create_test_document(font="Arial")
    file_path = "test_files/test.docx"
    doc.save(file_path)
    # Пример вызова проверки
    errors = formatting_check.check(doc, file_path)
    # Ожидаем ошибку, что используется неправильный шрифт
    assert any("Используется шрифт Arial, ожидается Times New Roman" in error for error in errors)

# Тест 2: Проверка правильного размера шрифта
def test_correct_font_size(formatting_check):
    doc = create_test_document(font_size=14)
    file_path = "test_files/test_correct_font_size.docx"
    doc.save(file_path)
    errors = formatting_check.check(doc, file_path, {"font_size": 14})
    assert "Размер шрифта 14 pt, ожидается 14 pt" not in errors

# Тест 3: Проверка неправильного размера шрифта
def test_wrong_font_size(formatting_check):
    doc = create_test_document(font_size=16)
    file_path = "test_files/test_wrong_font_size.docx"
    doc.save(file_path)
    errors = formatting_check.check(doc, file_path, {"font_size": 14})
    assert "Размер шрифта 16.0 pt, ожидается 14 pt" in errors

# Тест 4: Проверка выравнивания текста
def test_wrong_alignment(formatting_check):
    doc = create_test_document(alignment=WD_ALIGN_PARAGRAPH.LEFT)
    file_path = "test_files/test_wrong_alignment.docx"
    doc.save(file_path)
    errors = formatting_check.check(doc, file_path, {"alignment": WD_ALIGN_PARAGRAPH.CENTER})
    assert "Выравнивание параграфа должно быть по центру, найдено выравнивание по левому краю" in errors


# Тест 5: Проверка абзацного отступа
def test_wrong_first_line_indent(formatting_check):
    doc = create_test_document(first_line_indent=2)  # Задаем отступ первой строки
    file_path = "test_files/test_wrong_first_line_indent.docx"
    doc.save(file_path)
    errors = formatting_check.check(doc, file_path, {"first_line_indent": 1})
    assert "Отступ первой строки должен быть 1.0 inch, найден 2.0 inch" in errors


# Тест 6: Проверка правильного междустрочного интервала
def test_correct_line_spacing(formatting_check):
    doc = create_test_document()
    para = doc.add_paragraph("Another paragraph")
    para.paragraph_format.line_spacing = 1.5
    file_path = "test_files/test_correct_line_spacing.docx"
    doc.save(file_path)
    errors = formatting_check.check(doc, file_path, {"line_spacing": 1.5})
    assert "Междустрочный интервал" not in errors

# Тест 7: Проверка неправильного междустрочного интервала
def test_wrong_line_spacing(formatting_check):
    doc = create_test_document()
    para = doc.add_paragraph("Another paragraph")
    para.paragraph_format.line_spacing = 1.2
    file_path = "test_files/test_wrong_line_spacing.docx"
    doc.save(file_path)
    errors = formatting_check.check(doc, file_path, {"line_spacing": 1.5})
    assert "Междустрочный интервал 1.2, ожидается 1.5" in errors

# Тест 8: Проверка сноски с неправильным шрифтом
def test_footnote_with_wrong_font(formatting_check):
    doc = Document()
    para = doc.add_paragraph("Text in footnote")
    run = para.add_run("Footnote text")
    run.font.name = "Arial"
    para.add_run(" Text").font.name = "Arial"
    # Добавляем сноску
    footnote = doc.add_paragraph("Footnote paragraph")
    footnote_run = footnote.add_run("This is a footnote")
    footnote_run.font.name = "Arial"
    file_path = "test_files/test_footnote_with_wrong_font.docx"
    doc.save(file_path)
    errors = formatting_check.check(doc, file_path, {"font": "Times New Roman"})
    assert "Используется шрифт Arial, ожидается Times New Roman" in errors

# Тест 9: Проверка неправильного шрифта в таблице
def test_table_with_wrong_font(formatting_check):
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    cell = table.cell(0, 0)
    para = cell.add_paragraph("Text in table cell")
    run = para.add_run("Table cell text")
    run.font.name = "Arial"
    file_path = "test_files/test_table_with_wrong_font.docx"
    doc.save(file_path)
    errors = formatting_check.check(doc, file_path, {"font": "Times New Roman"})
    assert "Таблица 1, ячейка (1, 1), параграф 1: Используется шрифт Arial, ожидается Times New Roman" in errors

# Тест 10: Проверка пустого документа (не должно быть ошибок)
def test_empty_document(formatting_check):
    doc = Document()
    file_path = "test_files/test_empty_document.docx"
    doc.save(file_path)
    errors = formatting_check.check(doc, file_path, {"font": "Times New Roman"})
    assert len(errors) == 0
