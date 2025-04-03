import unittest
from docx import Document
from modules.appendices import AppendicesCheck  # Импортируйте нужный класс для проверки


class TestAppendicesCheck(unittest.TestCase):

    def setUp(self):
        # Создаем новый документ для каждого теста
        self.document = Document()

        # Добавляем приложения в документ для теста
        self.document.add_paragraph("Приложение A")  # Добавим приложение с буквой A
        self.document.add_paragraph("Тематический заголовок приложения A")
        self.document.add_paragraph("Ссылка на приложение А")
        self.document.add_paragraph("Оглавление")

        # Дополнительные параграфы и разрывы страниц для тестов
        self.document.add_paragraph("Текст документа")
        self.document.add_paragraph("Текст документа")
        self.document.add_paragraph("Текст документа")

    def test_appendix_numbering_alpha(self):
        # Создаем объект проверки с параметром нумерации как "alpha"
        checker = AppendicesCheck()
        errors = checker.check(self.document, {"appendix_number_style": "alpha"})

        # Проверяем, что ошибки по нумерации отсутствуют
        self.assertEqual(len(errors), 0)

    def test_appendix_numbering_numeric(self):
        # Создаем объект проверки с параметром нумерации как "numeric"
        checker = AppendicesCheck()
        errors = checker.check(self.document, {"appendix_number_style": "numeric"})

        # Проверяем, что ошибки по нумерации отсутствуют
        self.assertEqual(len(errors), 0)

    def test_check_for_appendices_in_toc(self):
        # Добавляем текст оглавления, который включает приложения
        self.document.add_paragraph("Оглавление")
        self.document.add_paragraph("Приложение A")

        checker = AppendicesCheck()
        errors = checker.check(self.document, {"appendix_number_style": "numeric"})

        # Проверяем, что ошибка о ненахождении приложения в оглавлении отсутствует
        self.assertEqual(len(errors), 0)

    def test_check_for_references_in_text(self):
        # Добавляем ссылку на приложение в тексте
        self.document.add_paragraph("Ссылка на приложение A")

        checker = AppendicesCheck()
        errors = checker.check(self.document, {"appendix_number_style": "alpha"})

        # Проверяем, что не возникает ошибок по отсутствующим ссылкам
        self.assertEqual(len(errors), 0)

    def test_invalid_appendix_number(self):
        # Изменим номер приложения, чтобы он был неверным
        self.document.paragraphs[0].text = "Приложение X"

        checker = AppendicesCheck()
        errors = checker.check(self.document, {"appendix_number_style": "numeric"})

        # Проверяем, что ошибка по неправильному номеру приложения будет найдена
        self.assertIn("Приложение (параграф 1): Неверный номер приложения, ожидается 1, получено 'X'", errors)

    def test_missing_page_break(self):
        # Убираем разрыв страницы перед первым приложением для проверки ошибки
        # Разрыв страницы для первого приложения
        self.document.paragraphs[1].add_run().add_break()  # Добавляем разрыв страницы перед приложением

        checker = AppendicesCheck()
        errors = checker.check(self.document, {"appendix_number_style": "numeric"})

        # Ожидаем, что ошибка о разрыве страницы будет найдена
        self.assertIn("Приложение 1 (параграф 1): Приложение должно начинаться с новой страницы", errors)


if __name__ == "__main__":
    unittest.main()
