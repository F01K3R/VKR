import unittest
from unittest.mock import MagicMock
from docx import Document
from modules.references import ReferencesCheck  # Импортируйте ваш класс проверки

class TestReferencesCheck(unittest.TestCase):

    def setUp(self):
        # Создаем реальный объект Document
        self.document = Document()
        # Добавляем параграфы в документ для тестов
        self.paragraphs = [self.document.add_paragraph() for _ in range(5)]

    def test_invalid_params_type(self):
        checker = ReferencesCheck()
        errors = checker.check(self.document, params=[])
        # Изменяем проверку ошибки для типа объекта
        self.assertTrue("Ошибка: params должен быть словарем" in errors[0] or "получено: <class 'list'>" in errors[0])

    def test_invalid_doc_type(self):
        checker = ReferencesCheck()
        errors = checker.check([], params={"standard": "ГОСТ Р 7.0.5-2008"})
        # Изменяем проверку ошибки, чтобы проверять тип объекта
        self.assertTrue(
            "Ошибка: doc должен быть объектом Document" in errors[0] or "получено: <class 'list'>" in errors[0])

    def test_missing_references_section(self):
        # Если раздел источников не найден
        self.paragraphs[0].text = "Введение"
        checker = ReferencesCheck()
        errors = checker.check(self.document, params={"standard": "ГОСТ Р 7.0.5-2008"})
        self.assertIn("No references section found", errors)

    def test_empty_references_section(self):
        # Если раздел источников пустой
        self.paragraphs[0].text = "Список литературы"
        self.paragraphs[1].text = ""
        checker = ReferencesCheck()
        errors = checker.check(self.document, params={"standard": "ГОСТ Р 7.0.5-2008"})
        self.assertIn("References section is empty", errors)

    def test_invalid_reference_format(self):
        # Если ссылка не соответствует стандарту
        self.paragraphs[1].text = "Неверная ссылка"
        checker = ReferencesCheck()
        errors = checker.check(self.document, params={"standard": "ГОСТ Р 7.0.5-2008"})
        self.assertIn("Неверный формат ссылки", errors)

    def test_correct_references(self):
        # Если ссылки соответствуют стандарту
        self.paragraphs[0].text = "Список источников"
        self.paragraphs[1].text = "1. Иванов И. И. Основы программирования. М.: Издательство, 2020. С. 100."
        self.paragraphs[2].text = "2. Петренко П. П. Теория алгоритмов. СПб.: Наука, 2018. С. 55."
        checker = ReferencesCheck()
        errors = checker.check(self.document, params={"standard": "ГОСТ Р 7.0.5-2008"})
        self.assertEqual(len(errors), 0)  # Ожидаем, что ошибок не будет

    def test_numbered_reference_check(self):
        # Проверка сквозной нумерации
        self.paragraphs[0].text = "Список источников"
        self.paragraphs[1].text = "1. Иванов И. И. Основы программирования. М.: Издательство, 2020. С. 100."
        self.paragraphs[2].text = "3. Петренко П. П. Теория алгоритмов. СПб.: Наука, 2018. С. 55."  # Пропущена ссылка с номером 2
        checker = ReferencesCheck()
        errors = checker.check(self.document, params={"standard": "ГОСТ Р 7.0.5-2008"})
        self.assertIn("Нарушение сквозной нумерации", errors)

    def test_alphabetical_order_check(self):
        # Проверка алфавитного порядка
        self.paragraphs[0].text = "Список литературы"
        self.paragraphs[1].text = "2. Петренко П. П. Теория алгоритмов. СПб.: Наука, 2018. С. 55."
        self.paragraphs[2].text = "1. Иванов И. И. Основы программирования. М.: Издательство, 2020. С. 100."
        checker = ReferencesCheck()
        errors = checker.check(self.document, params={"standard": "ГОСТ Р 7.0.5-2008"})
        self.assertIn("Иностранные источники", errors)

    def test_incorrect_citation_format(self):
        # Проверка неверного формата затекстовой ссылки
        self.paragraphs[0].text = "Текст с затекстовой ссылкой [Иванов И. И., 2020, с. 100]"
        checker = ReferencesCheck()
        errors = checker.check(self.document, params={"standard": "ГОСТ Р 7.0.5-2008"})
        self.assertIn("Неверный формат затекстовой ссылки", errors)

    def test_valid_citation(self):
        # Проверка правильного формата затекстовой ссылки
        self.paragraphs[0].text = "Текст с затекстовой ссылкой [Иванов И. И., 2020, с. 100]"
        checker = ReferencesCheck()
        errors = checker.check(self.document, params={"standard": "ГОСТ Р 7.0.5-2008"})
        self.assertEqual(len(errors), 0)  # Ожидаем, что ошибок не будет

if __name__ == "__main__":
    unittest.main()
