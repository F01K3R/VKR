import unittest
from docx import Document
from modules.tables import TablesCheck


class TestTablesCheck(unittest.TestCase):
    def setUp(self):
        # Создание документа для тестирования
        self.doc = Document()

        # Добавляем несколько параграфов для проверки
        self.doc.add_paragraph("Глава 1")
        self.doc.add_paragraph("Табл. 1 – Пример таблицы")
        table = self.doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Ячейка 1"
        table.cell(0, 1).text = "Ячейка 2"
        table.cell(1, 0).text = "Ячейка 3"
        table.cell(1, 1).text = "Ячейка 4"

        self.doc.add_paragraph("Таблица 1: Ссылка на таблицу")

    def test_table_caption_format(self):
        check = TablesCheck()
        errors = check.check(self.doc)

        # Проверяем наличие ошибок по формату заголовка таблицы
        self.assertNotIn("Неверный формат заголовка", errors)

    def test_empty_cell(self):
        # Создадим таблицу с пустыми ячейками
        self.doc.add_paragraph("Табл. 2 – Пример пустой ячейки")
        table = self.doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = ""
        table.cell(0, 1).text = "Ячейка 2"
        table.cell(1, 0).text = "Ячейка 3"
        table.cell(1, 1).text = ""

        check = TablesCheck()
        errors = check.check(self.doc)

        # Проверяем наличие ошибок по пустым ячейкам
        self.assertIn("Обнаружена пустая ячейка", errors)

    def test_table_numbering(self):
        # Проверка нумерации таблиц с параметром use_chapter_numbering
        self.doc.add_paragraph("Глава 2")
        self.doc.add_paragraph("Табл. 2 – Пример таблицы второй главы")
        table = self.doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Ячейка 1"
        table.cell(0, 1).text = "Ячейка 2"
        table.cell(1, 0).text = "Ячейка 3"
        table.cell(1, 1).text = "Ячейка 4"

        check = TablesCheck()
        errors = check.check(self.doc, params={"use_chapter_numbering": True})

        # Проверяем наличие ошибок по нумерации таблиц
        self.assertNotIn("Неверный номер таблицы", errors)

    def test_table_reference(self):
        # Создаем ссылку на таблицу в тексте
        self.doc.add_paragraph("Табл. 1 – Пример таблицы")
        table = self.doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Ячейка 1"
        table.cell(0, 1).text = "Ячейка 2"
        table.cell(1, 0).text = "Ячейка 3"
        table.cell(1, 1).text = "Ячейка 4"
        self.doc.add_paragraph("Таблица 1: Ссылка на таблицу")

        check = TablesCheck()
        errors = check.check(self.doc)

        # Проверяем наличие ошибок по ссылке на таблицу
        self.assertNotIn("Отсутствует ссылка на таблицу", errors)

    def test_caption_case(self):
        # Создание таблицы с неправильным регистром в заголовке
        self.doc.add_paragraph("Табл. 1 – пример таблицы с неправильным регистром")
        table = self.doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Ячейка 1"
        table.cell(0, 1).text = "Ячейка 2"
        table.cell(1, 0).text = "Ячейка 3"
        table.cell(1, 1).text = "Ячейка 4"

        check = TablesCheck()
        errors = check.check(self.doc)

        # Проверяем ошибку по регистру
        self.assertIn("Название 'пример таблицы с неправильным регистром'", errors)


if __name__ == '__main__':
    unittest.main()
