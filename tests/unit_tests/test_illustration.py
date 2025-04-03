import pytest
from docx import Document
from modules.illustrations import IllustrationsCheck


# Утилита для создания документа с рисунками
def create_test_document_with_figures(figures, use_chapter_numbering=False):
    doc = Document()

    # Добавляем главы и рисунки
    for i, (figure_num, caption) in enumerate(figures):
        doc.add_paragraph(f"ГЛАВА {i + 1}")
        doc.add_paragraph(f"Текст главы {i + 1}.")

        # Добавляем рисунок
        para = doc.add_paragraph()
        para.add_run("Рисунок").bold = True
        para.add_run(f" {figure_num}")
        para.add_run(" Рисунок показывающий что-то интересное.")

        # Добавляем подрисуночный текст
        doc.add_paragraph(f"Рис. {figure_num} – {caption}")

    return doc


# Тест на правильность нумерации рисунков с учетом главы
def test_correct_figure_numbering():
    figures = [
        ("1", "Описание первого рисунка"),
        ("2", "Описание второго рисунка"),
        ("3", "Описание третьего рисунка"),
    ]

    doc = create_test_document_with_figures(figures, use_chapter_numbering=True)
    checker = IllustrationsCheck()
    errors = checker.check(doc, params={"use_chapter_numbering": True})
    assert not errors, f"Ошибки: {errors}"


# Тест на отсутствие подрисуночного текста
def test_missing_caption():
    figures = [
        ("1", "Описание первого рисунка"),
        ("2", "Описание второго рисунка")
    ]

    doc = create_test_document_with_figures(figures, use_chapter_numbering=True)
    doc.paragraphs[1].text = "Рис. 1 – Название рисунка"  # Убираем подрисуночный текст
    checker = IllustrationsCheck()
    errors = checker.check(doc, params={"use_chapter_numbering": True})
    assert "Отсутствует подрисуночный текст" in errors, f"Ошибки: {errors}"


# Тест на проверку ссылок на рисунки
def test_missing_figure_references():
    figures = [
        ("1", "Описание первого рисунка"),
        ("2", "Описание второго рисунка")
    ]

    doc = create_test_document_with_figures(figures)
    checker = IllustrationsCheck()
    errors = checker.check(doc)
    assert "Отсутствует ссылка на рисунок" in errors, f"Ошибки: {errors}"


# Тест на наличие раздела "Список иллюстративного материала"
def test_missing_illustrations_list():
    figures = [
        ("1", "Описание первого рисунка"),
        ("2", "Описание второго рисунка")
    ]

    doc = create_test_document_with_figures(figures)
    doc.add_paragraph("Список литературы")
    checker = IllustrationsCheck()
    errors = checker.check(doc)
    assert "Отсутствует раздел 'Список иллюстративного материала'" in errors, f"Ошибки: {errors}"


# Тест на отсутствие ошибок при правильном оформлении
def test_correct_formatting():
    figures = [
        ("1", "Описание первого рисунка"),
        ("2", "Описание второго рисунка")
    ]

    doc = create_test_document_with_figures(figures)
    checker = IllustrationsCheck()
    errors = checker.check(doc)
    assert not errors, f"Ошибки: {errors}"
