import unittest
from unittest.mock import MagicMock
from docx.document import Document
from modules.structure import StructureCheck  # Замените на имя вашего модуля

class TestStructureCheck(unittest.TestCase):

    def setUp(self):
        # Инициализация документа для тестов
        self.document = MagicMock(spec=Document)
        self.document.paragraphs = [
            MagicMock(text="Оглавление", style=MagicMock(name="Heading 1")),
            MagicMock(text="Введение", style=MagicMock(name="Heading 1")),
            MagicMock(text="Основная часть", style=MagicMock(name="Heading 1")),
            MagicMock(text="Заключение", style=MagicMock(name="Heading 1")),
            MagicMock(text="Список литературы", style=MagicMock(name="Heading 1"))
        ]

    def test_invalid_params_type(self):
        # Проверка на неправильный тип параметра
        checker = StructureCheck()
        errors = checker.check(self.document, params=[])
        self.assertIn("Ошибка: params должен быть словарем", errors)

    def test_invalid_doc_type(self):
        # Проверка на неправильный тип документа
        checker = StructureCheck()
        errors = checker.check([], params={"require_headings": True})
        self.assertIn("Ошибка: document должен быть объектом Document", errors)

    def test_invalid_require_headings_type(self):
        # Проверка на неправильный тип require_headings
        checker = StructureCheck()
        errors = checker.check(self.document, params={"require_headings": "True"})
        self.assertIn("Ошибка: require_headings должен быть булевым значением", errors)

    def test_missing_required_section(self):
        # Проверка, если отсутствует обязательный раздел
        self.document.paragraphs = [MagicMock(text="Оглавление", style=MagicMock(name="Heading 1"))]
        checker = StructureCheck()
        errors = checker.check(self.document, params={"required_sections": ["Оглавление", "Введение"]})
        self.assertIn("Отсутствуют некоторые обязательные разделы", errors)

    def test_heading_format_check(self):
        # Проверка форматирования заголовков
        self.document.paragraphs = [
            MagicMock(text="Оглавление", style=MagicMock(name="Heading 1")),
            MagicMock(text="Введение", style=MagicMock(name="Heading 1")),
            MagicMock(text="Основная часть", style=MagicMock(name="Heading 1"))
        ]
        self.document.paragraphs[0].text = "Оглавление"
        self.document.paragraphs[1].text = "Введение"
        self.document.paragraphs[2].text = "Основная часть"
        checker = StructureCheck()
        errors = checker.check(self.document, params={"require_headings": True})
        self.assertNotIn("Заголовок", errors)

    def test_missing_headings(self):
        # Проверка, если заголовки отсутствуют
        self.document.paragraphs = [
            MagicMock(text="Оглавление"),
            MagicMock(text="Введение")
        ]
        checker = StructureCheck()
        errors = checker.check(self.document, params={"require_headings": True})
        self.assertIn("В документе отсутствуют заголовки", errors)

    def test_check_toc(self):
        # Проверка оглавления
        self.document.paragraphs = [
            MagicMock(text="Оглавление", style=MagicMock(name="Heading 1")),
            MagicMock(text="Глава 1. Введение", style=MagicMock(name="Heading 1")),
            MagicMock(text="Глава 2. Основная часть", style=MagicMock(name="Heading 1"))
        ]
        self.document.paragraphs[0].text = "Оглавление"
        self.document.paragraphs[1].text = "Глава 1. Введение"
        self.document.paragraphs[2].text = "Глава 2. Основная часть"
        checker = StructureCheck()
        errors = checker.check(self.document, params={"require_headings": True})
        self.assertNotIn("В оглавлении строка", errors)

    def test_missing_toc_section(self):
        # Проверка, если нет оглавления
        self.document.paragraphs = [MagicMock(text="Введение", style=MagicMock(name="Heading 1"))]
        checker = StructureCheck()
        errors = checker.check(self.document, params={"required_sections": ["Оглавление"]})
        self.assertIn("Отсутствуют некоторые обязательные разделы", errors)

    def test_check_section_order(self):
        # Проверка порядка разделов
        self.document.paragraphs = [
            MagicMock(text="Оглавление", style=MagicMock(name="Heading 1")),
            MagicMock(text="Введение", style=MagicMock(name="Heading 1")),
            MagicMock(text="Основная часть", style=MagicMock(name="Heading 1")),
            MagicMock(text="Заключение", style=MagicMock(name="Heading 1"))
        ]
        checker = StructureCheck()
        errors = checker.check(self.document, params={"required_sections": ["Оглавление", "Введение", "Основная часть", "Заключение"]})
        self.assertNotIn("Нарушение порядка разделов", errors)

    def test_invalid_heading_format(self):
        # Проверка неправильного формата заголовка (например, заголовок с точкой)
        self.document.paragraphs = [
            MagicMock(text="Глава 1: Введение", style=MagicMock(name="Heading 1"))
        ]
        checker = StructureCheck()
        errors = checker.check(self.document, params={"require_headings": True})
        self.assertIn("Заголовок 'Глава 1: Введение' (параграф 1) не должен заканчиваться точкой", errors)

if __name__ == "__main__":
    unittest.main()
