import unittest
from unittest.mock import MagicMock
from docx import Document
from modules.page_params import PageParamsCheck  # Импортируйте нужный класс для проверки

class TestPageParamsCheck(unittest.TestCase):

    def setUp(self):
        # Создаем объект документа
        self.document = MagicMock(spec=Document)

        # Создаем фиктивные секции документа для теста
        section_mock = MagicMock()
        section_mock.page_width.cm = 21.0  # ширина страницы
        section_mock.page_height.cm = 29.7  # высота страницы
        section_mock.left_margin.cm = 3.0  # левое поле
        section_mock.right_margin.cm = 1.0  # правое поле
        section_mock.top_margin.cm = 2.0  # верхнее поле
        section_mock.bottom_margin.cm = 2.0  # нижнее поле
        section_mock.header.paragraphs = [MagicMock(text="1", alignment=1)]  # Номер страницы в колонтитуле

        # Мокируем секции документа
        self.document.sections = [section_mock]

    def test_invalid_params_type(self):
        # Проверяем, что происходит, если параметры не переданы или передан неверный тип
        checker = PageParamsCheck()
        errors = checker.check(self.document, params=[])
        self.assertIn("Ошибка: params должен быть словарем", errors)

    def test_unsupported_page_size(self):
        # Проверяем, что произойдет при использовании неподдерживаемого размера страницы
        checker = PageParamsCheck()
        errors = checker.check(self.document, params={"page_size": "A5"})
        self.assertIn("Ошибка: неподдерживаемый размер страницы: A5", errors)

    def test_missing_margin_key(self):
        # Проверяем, что происходит, если в полях отсутствует ключ
        checker = PageParamsCheck()
        errors = checker.check(self.document, params={"margins": {"left": 3, "right": 1, "top": 2}})
        self.assertIn("Ошибка: в margins отсутствует ключ bottom", errors)

    def test_invalid_margin_value(self):
        # Проверяем, что происходит, если значение поля не является числом или меньше нуля
        checker = PageParamsCheck()
        errors = checker.check(self.document, params={"margins": {"left": 3, "right": -1, "top": 2, "bottom": 2}})
        self.assertIn("Ошибка: margins[right] должен быть неотрицательным числом", errors)

    def test_page_size_check(self):
        # Проверяем, что размер страницы соответствует ожидаемому
        checker = PageParamsCheck()
        errors = checker.check(self.document, params={"page_size": "A4", "margins": {"left": 3, "right": 1, "top": 2, "bottom": 2}})
        self.assertEqual(len(errors), 0)  # Ожидаем, что ошибок не будет

    def test_page_size_mismatch(self):
        # Проверяем ошибку, если размер страницы не соответствует ожидаемому
        self.document.sections[0].page_width.cm = 22.0  # Устанавливаем неверную ширину страницы
        checker = PageParamsCheck()
        errors = checker.check(self.document, params={"page_size": "A4"})
        self.assertIn("Секция 1: Размер страницы не соответствует ожидаемому (A4, портретная ориентация)", errors)

    def test_invalid_page_number(self):
        # Проверяем ошибку по нумерации страниц
        self.document.sections[0].header.paragraphs[0].text = "2"  # Неверный номер страницы
        checker = PageParamsCheck()
        errors = checker.check(self.document, params={"page_size": "A4"})
        self.assertIn("Вторая страница должна иметь номер 2", errors)

    def test_page_number_alignment(self):
        # Проверяем выравнивание номера страницы
        self.document.sections[0].header.paragraphs[0].alignment = 0  # Выравнивание по левому краю
        checker = PageParamsCheck()
        errors = checker.check(self.document, params={"page_size": "A4"})
        self.assertIn("Секция 1: Номер страницы должен быть выровнен по центру верхнего поля", errors)

    def test_correct_page_numbers(self):
        # Проверяем корректную нумерацию страниц
        self.document.sections[0].header.paragraphs[0].text = "1"  # Первая страница
        checker = PageParamsCheck()
        errors = checker.check(self.document, params={"page_size": "A4"})
        self.assertEqual(len(errors), 0)  # Ожидаем, что ошибок не будет

    def test_correct_orientation_portrait(self):
        # Проверяем, что ориентация страницы портретная
        checker = PageParamsCheck()
        errors = checker.check(self.document, params={"page_size": "A4"})
        self.assertEqual(len(errors), 0)  # Ожидаем, что ошибок не будет

    def test_incorrect_orientation_landscape(self):
        # Проверяем ошибку, если ориентация страницы неверная
        self.document.sections[0].page_width.cm = 29.7  # Сделаем страницу альбомной ориентации
        checker = PageParamsCheck()
        errors = checker.check(self.document, params={"page_size": "A4"})
        self.assertIn("Секция 1: Ориентация страницы альбомная, ожидается портретная", errors)

if __name__ == "__main__":
    unittest.main()
