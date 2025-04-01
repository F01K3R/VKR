import re
import logging
from docx.document import Document
from docx.oxml.ns import qn
from .base import CheckModule

logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

class StructureCheck(CheckModule):
    def check(self, document, params=None):
        # Проверка входных параметров
        if params is None:
            params = {}
        if not isinstance(params, dict):
            return [f"Ошибка: params должен быть словарем, получено: {type(params)}"]
        if not isinstance(document, Document):
            return [f"Ошибка: document должен быть объектом Document, получено: {type(document)}"]

        require_headings = params.get("require_headings", True)
        if not isinstance(require_headings, bool):
            return [f"Ошибка: require_headings должен быть булевым значением, получено: {type(require_headings)}"]

        required_sections = params.get("required_sections", [
            "Оглавление",
            "Введение",
            "Основная часть",
            "Заключение",
            "Список литературы"
        ])
        if not isinstance(required_sections, list):
            return [f"Ошибка: required_sections должен быть списком, получено: {type(required_sections)}"]
        if not all(isinstance(s, str) for s in required_sections):
            return [f"Ошибка: все элементы required_sections должны быть строками"]

        errors = []

        # Подготовка для поиска
        required_sections_lower = [s.lower() for s in required_sections]
        section_patterns = [
            re.compile(rf"^(?:\d+\.\s*)?{re.escape(section_lower)}(?:\.|:)?\s*$")
            for section_lower in required_sections_lower
        ]

        # Проверка наличия заголовков и разделов
        headings = []
        heading_levels = {}
        found_sections = {}
        paragraphs = []
        toc_content = []  # Для хранения содержимого оглавления
        in_toc = False

        try:
            for i, para in enumerate(document.paragraphs):
                text = para.text.strip() if para.text else ""
                paragraphs.append((i, para, text))

                # Проверка заголовков
                try:
                    if para.style and para.style.name and para.style.name.startswith('Heading'):
                        headings.append(para)
                        text_upper = text.upper()
                        level = int(para.style.name.split()[-1])
                        heading_levels[text] = level

                        # Проверка оформления заголовков
                        if text_upper != text:
                            errors.append(f"Заголовок '{text}' (параграф {i+1}) должен быть в верхнем регистре")
                        if text.endswith('.'):
                            errors.append(f"Заголовок '{text}' (параграф {i+1}) не должен заканчиваться точкой")
                        if '-' in text and not re.search(r'\b(ОАО|АО|ООО|ЗАО)\b', text):
                            errors.append(f"Заголовок '{text}' (параграф {i+1}) содержит недопустимый перенос или сокращение")
                        # Проверка интервала после заголовка (должно быть 1.5)
                        if i + 1 < len(document.paragraphs):
                            next_para = document.paragraphs[i + 1]
                            if next_para.paragraph_format.line_spacing != 1.5:
                                errors.append(f"После заголовка '{text}' (параграф {i+1}) интервал должен быть 1.5, текущий: {next_para.paragraph_format.line_spacing}")
                        # Проверка, начинается ли глава с новой страницы
                        if "ГЛАВА" in text_upper:
                            prev_elements = para._element.xpath('preceding-sibling::*')
                            has_page_break = any(elem.tag.endswith('br') and elem.get(qn('w:type')) == 'page' for elem in prev_elements)
                            if not has_page_break:
                                errors.append(f"Глава '{text}' (параграф {i+1}) должна начинаться с новой страницы")
                except Exception as e:
                    return [f"Ошибка при проверке стиля параграфа: {str(e)}"]

                # Проверка разделов
                text_lower = text.lower()
                for section, pattern in zip(required_sections, section_patterns):
                    if pattern.match(text_lower):
                        found_sections[section] = i
                        # Проверка уровня заголовка
                        if text in heading_levels and heading_levels[text] != 1:
                            errors.append(f"Раздел '{section}' имеет стиль 'Heading {heading_levels[text]}', ожидается 'Heading 1'")
                        # Проверка, является ли это оглавлением
                        if section.lower() == "оглавление":
                            in_toc = True
                        break
                    elif in_toc and text:
                        toc_content.append((text, i))
                    elif in_toc and re.match(r"^(Введение|Список сокращений|Список терминов)", text):
                        in_toc = False

        except Exception as e:
            return [f"Ошибка при доступе к параграфам документа: {str(e)}"]

        # Проверка наличия заголовков
        if require_headings and not headings:
            errors.append("В документе отсутствуют заголовки (стиль 'Heading')")

        # Проверка отсутствующих разделов
        missing_sections = [s for s in required_sections if s not in found_sections]
        if missing_sections:
            logger.debug(f"Отсутствуют разделы: {', '.join(missing_sections)}")
            errors.append("Отсутствуют некоторые обязательные разделы")

        # Проверка порядка разделов
        if len(found_sections) > 1:
            found_sections_list = [(section, pos) for section, pos in found_sections.items()]
            found_sections_list.sort(key=lambda x: x[1])
            found_sections_order = [section for section, _ in found_sections_list]
            expected_order = [s for s in required_sections if s in found_sections_order]
            if found_sections_order != expected_order:
                errors.append(f"Нарушение порядка разделов: текущий порядок {found_sections_order}, ожидается {expected_order}")

        # Проверка оформления оглавления
        if "Оглавление" in found_sections:
            allowed_abbreviations = r'\b(ОАО|АО|ООО|ЗАО)\b'
            for toc_line, idx in toc_content:
                # Проверка, что заголовки в верхнем регистре
                if not toc_line.isupper():
                    errors.append(f"В оглавлении строка '{toc_line}' (параграф {idx+1}) должна быть в верхнем регистре")
                # Проверка на отсутствие сокращений
                if '-' in toc_line and not re.search(allowed_abbreviations, toc_line):
                    errors.append(f"В оглавлении строка '{toc_line}' (параграф {idx+1}) содержит недопустимые сокращения")
                # Проверка отточия перед номером страницы
                if not re.match(r'.*\.\.\.\s*\d+$', toc_line):
                    errors.append(f"В оглавлении строка '{toc_line}' (параграф {idx+1}) должна заканчиваться отточием и номером страницы")
                # Проверка совпадения заголовков
                found = False
                for _, para, para_text in paragraphs:
                    if para_text.strip() == toc_line.split('...')[0].strip():
                        found = True
                        break
                if not found:
                    errors.append(f"В оглавлении строка '{toc_line}' (параграф {idx+1}) не соответствует ни одному заголовку в тексте")

        return errors