import re
import logging
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .base import CheckModule

logging.basicConfig(filename='processing.log', level=logging.DEBUG)
logger = logging.getLogger(__name__)

class IllustrationsCheck(CheckModule):
    # Регулярные выражения для поиска ссылок на рисунки в тексте
    FIGURE_REF_PATTERN = re.compile(r'(?:рисунок|рис\.)\s+(\d+(?:\.\d+)?)', re.IGNORECASE)
    # Регулярное выражение для подрисуночного текста
    FIGURE_CAPTION_PATTERN = re.compile(r'^Рис\.\s+(\d+(?:\.\d+)?)\s*–\s*(.+)$')
    # Регулярное выражение для поиска раздела "Список иллюстративного материала"
    ILLUSTRATIONS_LIST_PATTERN = re.compile(r'^(?:Список\s+иллюстративного\s+материала|Список\s+рисунков)$', re.IGNORECASE)

    def check(self, document, params=None):
        """
        Проверяет оформление иллюстраций в документе согласно требованиям.

        Args:
            document (Document): Объект документа .docx.
            params (dict, optional): Параметры проверки.

        Returns:
            list: Список ошибок, если они есть.
        """
        # Проверка входных параметров
        if params is None:
            params = {}
        if not isinstance(params, dict):
            return [f"Ошибка: params должен быть словарем, получено: {type(params)}"]
        if not isinstance(document, Document):
            return [f"Ошибка: document должен быть объектом Document, получено: {type(document)}"]

        # Параметры проверки
        use_chapter_numbering = params.get("use_chapter_numbering", False)  # Использовать нумерацию в пределах главы (например, 1.1)

        errors = []
        figure_positions = []  # Список кортежей (номер рисунка, индекс параграфа подписи)
        figure_references = []  # Список кортежей (номер рисунка, индекс параграфа ссылки)
        chapter_numbers = {}  # Словарь для хранения номеров глав (например, "1" для "ГЛАВА 1")
        in_appendices = False  # Флаг для проверки, находятся ли иллюстрации в приложении
        illustrations_list_idx = None  # Индекс раздела "Список иллюстративного материала"
        toc_content = []  # Содержимое оглавления
        in_toc = False  # Флаг для определения, находимся ли в оглавлении

        # Сбор всех параграфов и поиск иллюстраций
        paragraphs = []
        current_chapter = "0"  # По умолчанию, если глав нет
        for i, para in enumerate(document.paragraphs):
            text = para.text.strip() if para.text else ""
            paragraphs.append((i, para, text))

            # Определяем текущую главу для нумерации рисунков
            if text.upper().startswith("ГЛАВА"):
                match = re.match(r'ГЛАВА\s+(\d+)', text, re.IGNORECASE)
                if match:
                    current_chapter = match.group(1)
            chapter_numbers[i] = current_chapter

            # Определяем, находимся ли в приложении
            if text.lower().startswith("приложение"):
                in_appendices = True

            # Ищем раздел "Список иллюстративного материала"
            if self.ILLUSTRATIONS_LIST_PATTERN.match(text):
                illustrations_list_idx = i

            # Ищем оглавление
            if text.lower() == "оглавление":
                in_toc = True
            elif in_toc and text:
                toc_content.append(text)
            elif in_toc and re.match(r"^(Введение|Список сокращений|Список терминов)", text):
                in_toc = False

            # Ищем ссылки на рисунки в тексте
            matches = self.FIGURE_REF_PATTERN.findall(text)
            for figure_num in matches:
                figure_references.append((figure_num, i))

        # Проверяем иллюстрации
        expected_figure_num = 1
        figures_found = 0
        for para_idx, para in enumerate(document.paragraphs):
            # Ищем рисунки в параграфе (через <w:drawing> или <w:pict>)
            has_drawing = any(run._element.xpath('.//w:drawing') or run._element.xpath('.//w:pict') for run in para.runs)
            if not has_drawing:
                continue
            figures_found += 1

            # Ищем подрисуночный текст (следующий параграф после рисунка)
            caption_idx = para_idx + 1
            if caption_idx >= len(document.paragraphs):
                errors.append(f"Рисунок {figures_found}: Отсутствует подрисуночный текст после рисунка (параграф {para_idx+1})")
                continue
            caption_para = document.paragraphs[caption_idx]
            caption_text = caption_para.text.strip()
            match = self.FIGURE_CAPTION_PATTERN.match(caption_text)
            if not match:
                errors.append(f"Рисунок {figures_found}: Неверный формат подрисуночного текста (параграф {caption_idx+1}): '{caption_text}', ожидается 'Рис. N – Название'")
                continue

            # Извлекаем номер и название
            figure_num, figure_title = match.groups()
            figure_positions.append((figure_num, caption_idx))

            # Проверка нумерации
            if use_chapter_numbering:
                chapter = chapter_numbers.get(caption_idx, "0")
                expected_caption = f"{chapter}.{expected_figure_num}"
            else:
                expected_caption = str(expected_figure_num)
            if figure_num != expected_caption:
                errors.append(f"Рисунок {figures_found}: Неверный номер рисунка, ожидается {expected_caption}, получено {figure_num}")
            expected_figure_num += 1

            # Проверка оформления подрисуночного текста
            if not figure_title[0].isupper():
                errors.append(f"Рисунок {figure_num}: Название '{figure_title}' должно начинаться с прописной буквы (параграф {caption_idx+1})")
            if figure_title.endswith('.'):
                errors.append(f"Рисунок {figure_num}: Название '{figure_title}' не должно заканчиваться точкой (параграф {caption_idx+1})")
            if caption_para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                errors.append(f"Рисунок {figure_num}: Подрисуночный текст '{caption_text}' должен быть выровнен по центру (параграф {caption_idx+1})")

        # Проверка ссылок на рисунки
        referenced_figures = set(ref[0] for ref in figure_references)
        for figure_num, _ in figure_positions:
            if figure_num not in referenced_figures:
                errors.append(f"Рисунок {figure_num}: Отсутствует ссылка на рисунок в тексте")

        # Проверка раздела "Список иллюстративного материала"
        if not in_appendices and figures_found > 0:
            if illustrations_list_idx is None:
                errors.append("Отсутствует раздел 'Список иллюстративного материала' после списка литературы, хотя иллюстрации присутствуют в тексте")
            else:
                # Проверяем, что раздел включён в оглавление
                illustrations_list_title = document.paragraphs[illustrations_list_idx].text.strip()
                if not any(illustrations_list_title.upper() in toc_line.upper() for toc_line in toc_content):
                    errors.append("Раздел 'Список иллюстративного материала' не включён в оглавление")

        return errors