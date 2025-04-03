import re
import logging
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from modules.base import CheckModule

logging.basicConfig(filename='processing.log', level=logging.DEBUG)
logger = logging.getLogger(__name__)

class TablesCheck(CheckModule):
    # Регулярные выражения для поиска ссылок на таблицы в тексте
    TABLE_REF_PATTERN = re.compile(r'(?:таблица|табл\.)\s+(\d+(?:\.\d+)?)', re.IGNORECASE)
    # Регулярное выражение для заголовков таблиц
    TABLE_CAPTION_PATTERN = re.compile(r'^Табл\.\s+(\d+(?:\.\d+)?)\s*–\s*(.+)$')

    def check(self, document, params=None):
        """
        Проверяет оформление таблиц в документе согласно требованиям.

        Args:
            document (Document): Объект документа .docx.
            params (dict, optional): Параметры проверки.

        Returns:
            list: Список ошибок, если они есть.
        """
        # Проверка входных параметров
        if params is None:
            params = {}
        if not isinstance(params, dict):
            return [f"Ошибка: params должен быть словарем, получено: {type(params)}"]
        if not isinstance(document, Document):
            return [f"Ошибка: document должен быть объектом Document, получено: {type(document)}"]

        # Параметры проверки
        use_chapter_numbering = params.get("use_chapter_numbering", False)

        errors = []
        table_positions = []  # Список кортежей (номер таблицы, индекс таблицы, индекс параграфа подписи)
        table_references = []  # Список кортежей (номер таблицы, индекс параграфа ссылки)
        chapter_numbers = {}  # Словарь для хранения номеров глав
        current_chapter = "0"  # По умолчанию, если глав нет

        # Сбор всех параграфов и определение текущей главы
        for i, para in enumerate(document.paragraphs):
            text = para.text.strip() if para.text else ""
            # Определяем текущую главу для нумерации таблиц
            if text.upper().startswith("ГЛАВА"):
                match = re.match(r'ГЛАВА\s+(\d+)', text, re.IGNORECASE)
                if match:
                    current_chapter = match.group(1)
            chapter_numbers[i] = current_chapter

            # Ищем ссылки на таблицы в тексте
            matches = self.TABLE_REF_PATTERN.findall(text)
            for table_num in matches:
                table_references.append((table_num, i))

        # Проверяем таблицы
        expected_table_num = 1
        for table_idx, table in enumerate(document.tables):
            # Ищем заголовок таблицы (предшествующий параграф)
            caption_idx = None
            for i, para in enumerate(document.paragraphs):
                text = para.text.strip()
                match = self.TABLE_CAPTION_PATTERN.match(text)
                if match:
                    # Проверяем, относится ли этот заголовок к текущей таблице
                    # Для этого ищем следующую таблицу после параграфа
                    next_table_idx = None
                    for j, t in enumerate(document.tables):
                        if t._element.xpath(f'preceding::w:p[{i+1}]'):
                            next_table_idx = j
                            break
                    if next_table_idx == table_idx:
                        caption_idx = i
                        break

            if caption_idx is None:
                errors.append(f"Таблица {table_idx+1}: Отсутствует заголовок перед таблицей")
                continue

            caption_para = document.paragraphs[caption_idx]
            caption_text = caption_para.text.strip()
            match = self.TABLE_CAPTION_PATTERN.match(caption_text)
            if not match:
                errors.append(f"Таблица {table_idx+1}: Неверный формат заголовка (параграф {caption_idx+1}): '{caption_text}', ожидается 'Табл. N – Название'")
                continue

            # Извлекаем номер и название
            table_num, table_title = match.groups()
            table_positions.append((table_num, table_idx, caption_idx))

            # Проверка нумерации
            if use_chapter_numbering:
                chapter = chapter_numbers.get(caption_idx, "0")
                expected_caption = f"{chapter}.{expected_table_num}"
            else:
                expected_caption = str(expected_table_num)
            if table_num != expected_caption:
                errors.append(f"Таблица {table_idx+1}: Неверный номер таблицы, ожидается {expected_caption}, получено {table_num}")
            expected_table_num += 1

            # Проверка оформления заголовка
            if not table_title[0].isupper():
                errors.append(f"Таблица {table_num}: Название '{table_title}' должно начинаться с прописной буквы (параграф {caption_idx+1})")
            if table_title.endswith('.'):
                errors.append(f"Таблица {table_num}: Название '{table_title}' не должно заканчиваться точкой (параграф {caption_idx+1})")
            if caption_para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                errors.append(f"Таблица {table_num}: Заголовок '{caption_text}' должен быть выровнен по центру (параграф {caption_idx+1})")

            # Проверка содержимого таблицы
            for row in table.rows:
                for cell in row.cells:
                    if not cell.text.strip():
                        errors.append(f"Таблица {table_num}: Обнаружена пустая ячейка (таблица {table_idx+1})")

        # Проверка ссылок на таблицы
        referenced_tables = set(ref[0] for ref in table_references)
        for table_num, table_idx, caption_idx in table_positions:
            if table_num not in referenced_tables:
                errors.append(f"Таблица {table_num}: Отсутствует ссылка на таблицу в тексте (таблица {table_idx+1}, заголовок в параграфе {caption_idx+1})")

        return errors