import logging
from docx.document import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from modules.base import CheckModule
from utils.xml_utils import extract_xml

logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)


class FormattingCheck(CheckModule):
    def check(self, document, file_path, params=None):
        # Проверка входных параметров
        if params is None:
            params = {}
        if not isinstance(params, dict):
            return [f"Ошибка: params должен быть словарем, получено: {type(params)}"]
        if not isinstance(document, Document):
            return [f"Ошибка: document должен быть объектом Document, получено: {type(document)}"]
        if not isinstance(file_path, str):
            return [f"Ошибка: file_path должен быть строкой, получено: {type(file_path)}"]

        expected_font = params.get("font", "Times New Roman")
        expected_font_size = params.get("font_size", 14)  # в pt
        expected_line_spacing = params.get("line_spacing", 1.5)
        expected_alignment = params.get("alignment", "justify")
        expected_indent = params.get("first_line_indent", 1.25)  # в см

        errors = []

        # Извлечение XML для анализа форматирования
        try:
            document_xml, styles_xml = extract_xml(file_path)
        except Exception as e:
            return [f"Ошибка при извлечении XML: {str(e)}"]

        # Проверка стилей в styles.xml
        style_fonts = {}
        style_sizes = {}
        try:
            for style in styles_xml.findall(".//w:style", namespaces={
                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}):
                style_id = style.get(qn("w:styleId"))
                font = style.find(".//w:rFonts",
                                  namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
                size = style.find(".//w:sz",
                                  namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
                if font is not None:
                    font_name = font.get(qn("w:ascii"))
                    if font_name:
                        style_fonts[style_id] = font_name
                if size is not None:
                    size_val = size.get(qn("w:val"))
                    if size_val:
                        style_sizes[style_id] = int(size_val) / 2  # Размер в half-points, переводим в pt
        except Exception as e:
            errors.append(f"Ошибка при анализе стилей: {str(e)}")

        # Проверка форматирования параграфов
        for i, para in enumerate(document.paragraphs):
            try:
                # Пропускаем пустые параграфы
                if not para.text.strip():
                    continue

                # Проверка стиля параграфа
                style_id = para.style.style_id if para.style else None
                is_heading = style_id and style_id.startswith("Heading")

                # Проверка шрифта
                font_name = None
                run_fonts = set()
                for run in para.runs:
                    if run.font.name:
                        run_fonts.add(run.font.name)
                    elif style_id and style_id in style_fonts:
                        run_fonts.add(style_fonts[style_id])
                if run_fonts:
                    font_name = run_fonts.pop() if len(run_fonts) == 1 else None
                    if font_name and font_name != expected_font:
                        # Исключение: допускается использование других шрифтов для акцентирования
                        if not (is_heading or "формул" in para.text.lower() or "теорем" in para.text.lower()):
                            errors.append(
                                f"Параграф {i + 1}: Используется шрифт {font_name}, ожидается {expected_font}")

                # Проверка цвета шрифта (должен быть чёрным)
                for run in para.runs:
                    if run.font.color and run.font.color.rgb != (0, 0, 0):
                        errors.append(f"Параграф {i + 1}: Цвет шрифта должен быть чёрным, обнаружен другой цвет")

                # Проверка размера шрифта
                font_size = None
                run_sizes = set()
                for run in para.runs:
                    if run.font.size:
                        run_sizes.add(run.font.size.pt)
                    elif style_id and style_id in style_sizes:
                        run_sizes.add(style_sizes[style_id])
                if run_sizes:
                    font_size = run_sizes.pop() if len(run_sizes) == 1 else None
                    expected_size = 12 if R"^сноск" in para.text.lower() or r"^таблиц" in para.text.lower() or r"^приложени" in para.text.lower() or r"^рис" in para.text.lower() else expected_font_size
                    if font_size and font_size != expected_size:
                        errors.append(f"Параграф {i + 1}: Размер шрифта {font_size} pt, ожидается {expected_size} pt")

                # Проверка выравнивания
                alignment = para.alignment
                if alignment is not None:
                    if is_heading:
                        if alignment != WD_ALIGN_PARAGRAPH.CENTER:
                            errors.append(
                                f"Параграф {i + 1}: Заголовок должен быть выровнен по центру, текущее выравнивание: {alignment}")
                    else:
                        if expected_alignment == "justify" and alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
                            errors.append(
                                f"Параграф {i + 1}: Выравнивание должно быть по ширине, текущее выравнивание: {alignment}")
                        elif expected_alignment == "left" and alignment != WD_ALIGN_PARAGRAPH.LEFT:
                            errors.append(
                                f"Параграф {i + 1}: Выравнивание должно быть по левому краю, текущее выравнивание: {alignment}")

                # Проверка междустрочного интервала
                line_spacing = para.paragraph_format.line_spacing
                if line_spacing is not None and line_spacing != expected_line_spacing:
                    # Проверка интервала после заголовков уже есть в structure.py, здесь проверяем только основной текст
                    if not is_heading:
                        errors.append(
                            f"Параграф {i + 1}: Междустрочный интервал {line_spacing}, ожидается {expected_line_spacing}")

                # Проверка абзацного отступа
                first_line_indent = para.paragraph_format.first_line_indent
                if first_line_indent is not None:
                    indent_cm = first_line_indent.cm if first_line_indent else 0
                    if abs(indent_cm - expected_indent) > 0.01:  # Допуск 0.01 см
                        errors.append(
                            f"Параграф {i + 1}: Абзацный отступ {indent_cm:.2f} см, ожидается {expected_indent} см")

                # Проверка отсутствия дополнительных отступов (кроме абзацного)
                left_indent = para.paragraph_format.left_indent.cm if para.paragraph_format.left_indent else 0
                right_indent = para.paragraph_format.right_indent.cm if para.paragraph_format.right_indent else 0
                if left_indent != 0 or right_indent != 0:
                    errors.append(
                        f"Параграф {i + 1}: Дополнительные отступы слева ({left_indent} см) или справа ({right_indent} см) не допускаются")

            except Exception as e:
                errors.append(f"Параграф {i + 1}: Ошибка при проверке форматирования: {str(e)}")

        # Проверка форматирования в таблицах
        for table_idx, table in enumerate(document.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, para in enumerate(cell.paragraphs):
                        try:
                            if not para.text.strip():
                                continue
                            # Проверка шрифта в таблицах
                            run_fonts = set()
                            for run in para.runs:
                                if run.font.name:
                                    run_fonts.add(run.font.name)
                            if run_fonts:
                                font_name = run_fonts.pop() if len(run_fonts) == 1 else None
                                if font_name and font_name != expected_font:
                                    errors.append(
                                        f"Таблица {table_idx + 1}, ячейка ({row_idx + 1}, {cell_idx + 1}), параграф {para_idx + 1}: Используется шрифт {font_name}, ожидается {expected_font}")
                            # Проверка размера шрифта в таблицах (должен быть 12 pt)
                            run_sizes = set()
                            for run in para.runs:
                                if run.font.size:
                                    run_sizes.add(run.font.size.pt)
                            if run_sizes:
                                font_size = run_sizes.pop() if len(run_sizes) == 1 else None
                                if font_size and font_size != 12:
                                    errors.append(
                                        f"Таблица {table_idx + 1}, ячейка ({row_idx + 1}, {cell_idx + 1}), параграф {para_idx + 1}: Размер шрифта {font_size} pt, ожидается 12 pt")
                        except Exception as e:
                            errors.append(
                                f"Таблица {table_idx + 1}, ячейка ({row_idx + 1}, {cell_idx + 1}), параграф {para_idx + 1}: Ошибка при проверке форматирования: {str(e)}")

        # Проверка форматирования в сносках
        try:
            # Получаем часть документа со сносками через отношения
            rels = document.part.rels
            footnotes_part = None
            for rel in rels.values():
                if rel.reltype == "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes":
                    footnotes_part = rel.target_part
                    break

            if footnotes_part:
                # Извлекаем XML из footnotes_part
                import xml.etree.ElementTree as ET
                footnotes_xml = ET.fromstring(footnotes_part.blob)
                footnotes = footnotes_xml.findall(".//w:footnote", namespaces={
                    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})

                for footnote_idx, footnote in enumerate(footnotes):
                    # Пропускаем служебные сноски (например, footnote с id="-1" или "0")
                    footnote_id = footnote.get(qn("w:id"))
                    if footnote_id in ("-1", "0"):
                        continue

                    # Извлекаем параграфы из сноски
                    for para_idx, para in enumerate(footnote.findall(".//w:p", namespaces={
                        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})):
                        try:
                            # Извлекаем текст параграфа
                            para_text = "".join([t.text for t in para.findall(".//w:t", namespaces={
                                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}) if t.text])
                            if not para_text.strip():
                                continue

                            # Проверка шрифта в сносках
                            run_fonts = set()
                            for run in para.findall(".//w:r", namespaces={
                                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}):
                                font = run.find(".//w:rFonts", namespaces={
                                    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
                                if font is not None:
                                    font_name = font.get(qn("w:ascii"))
                                    if font_name:
                                        run_fonts.add(font_name)
                            if run_fonts:
                                font_name = run_fonts.pop() if len(run_fonts) == 1 else None
                                if font_name and font_name != expected_font:
                                    errors.append(
                                        f"Сноска {footnote_idx + 1}, параграф {para_idx + 1}: Используется шрифт {font_name}, ожидается {expected_font}")

                            # Проверка размера шрифта в сносках (должен быть 12 pt)
                            run_sizes = set()
                            for run in para.findall(".//w:r", namespaces={
                                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}):
                                size = run.find(".//w:sz", namespaces={
                                    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
                                if size is not None:
                                    size_val = size.get(qn("w:val"))
                                    if size_val:
                                        run_sizes.add(int(size_val) / 2)  # Размер в half-points, переводим в pt
                            if run_sizes:
                                font_size = run_sizes.pop() if len(run_sizes) == 1 else None
                                if font_size and font_size != 12:
                                    errors.append(
                                        f"Сноска {footnote_idx + 1}, параграф {para_idx + 1}: Размер шрифта {font_size} pt, ожидается 12 pt")
                        except Exception as e:
                            errors.append(
                                f"Сноска {footnote_idx + 1}, параграф {para_idx + 1}: Ошибка при проверке форматирования: {str(e)}")
            else:
                logger.debug("Сноски в документе отсутствуют.")
        except Exception as e:
            errors.append(f"Ошибка при доступе к сноскам: {str(e)}")

        # Проверка форматирования в приложениях (предполагаем, что приложения начинаются после раздела "Приложения")
        in_appendices = False
        for i, para in enumerate(document.paragraphs):
            if para.text.strip().lower().startswith("приложение"):
                in_appendices = True
            if in_appendices and para.text.strip():
                try:
                    # Проверка шрифта в приложениях
                    run_fonts = set()
                    for run in para.runs:
                        if run.font.name:
                            run_fonts.add(run.font.name)
                    if run_fonts:
                        font_name = run_fonts.pop() if len(run_fonts) == 1 else None
                        if font_name and font_name != expected_font:
                            errors.append(
                                f"Приложение, параграф {i + 1}: Используется шрифт {font_name}, ожидается {expected_font}")
                    # Проверка размера шрифта в приложениях (должен быть 12 pt)
                    run_sizes = set()
                    for run in para.runs:
                        if run.font.size:
                            run_sizes.add(run.font.size.pt)
                    if run_sizes:
                        font_size = run_sizes.pop() if len(run_sizes) == 1 else None
                        if font_size and font_size != 12:
                            errors.append(
                                f"Приложение, параграф {i + 1}: Размер шрифта {font_size} pt, ожидается 12 pt")
                except Exception as e:
                    errors.append(f"Приложение, параграф {i + 1}: Ошибка при проверке форматирования: {str(e)}")

        # Добавляем примечание о допустимом использовании других шрифтов
        if not any("допускается использование других шрифтов" in error for error in errors):
            errors.append(
                "Примечание: Допускается использование шрифтов разной гарнитуры для акцентирования терминов, формул, теорем")

        return errors