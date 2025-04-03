import re
import logging
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from .base import CheckModule

logging.basicConfig(filename='processing.log', level=logging.DEBUG)
logger = logging.getLogger(__name__)

class AppendicesCheck(CheckModule):
    # Регулярные выражения для поиска приложений
    APPENDIX_HEADER_PATTERN = re.compile(r'^Приложение\s+([А-Я]|\d+)$')
    # Регулярное выражение для поиска ссылок на приложения в тексте
    APPENDIX_REF_PATTERN = re.compile(r'(?:приложение|прил\.)\s+([А-Я]|\d+)', re.IGNORECASE)
    # Регулярное выражение для поиска списка литературы
    REFERENCES_HEADER_PATTERN = re.compile(r'^(?:Список\s+(?:источников|литературы)|Литература|Bibliography|References|Список\s+использованных\s+источников)$', re.IGNORECASE)
    # Регулярное выражение для поиска списка иллюстративного материала
    ILLUSTRATIONS_LIST_PATTERN = re.compile(r'^(?:Список\s+иллюстративного\s+материала|Список\s+рисунков)$', re.IGNORECASE)

    def check(self, document, params=None):
        """
        Проверяет оформление приложений в документе согласно требованиям.

        Args:
            document (Document): Объект документа .docx.
            params (dict, optional): Параметры проверки.

        Returns:
            list: Список ошибок, если они есть.
        """
        # Проверка входных параметров
        if params is None:
            params = {}
        if not isinstance(params, dict):
            return [f"Ошибка: params должен быть словарем, получено: {type(params)}"]
        if not isinstance(document, Document):
            return [f"Ошибка: document должен быть объектом Document, получено: {type(document)}"]

        # Параметры проверки
        appendix_number_style = params.get("appendix_number_style", "numeric")  # "numeric" (1, 2, 3) или "alpha" (А, Б, В)

        errors = []
        appendix_positions = []  # Список кортежей (номер приложения, индекс параграфа заголовка)
        appendix_references = []  # Список кортежей (номер приложения, индекс параграфа ссылки)
        references_idx = None  # Индекс раздела "Список литературы"
        illustrations_list_idx = None  # Индекс раздела "Список иллюстративного материала"
        toc_content = []  # Содержимое оглавления
        in_toc = False  # Флаг для определения, находимся ли в оглавлении

        # Сбор всех параграфов
        paragraphs = []
        for i, para in enumerate(document.paragraphs):
            text = para.text.strip() if para.text else ""
            paragraphs.append((i, para, text))

            # Ищем раздел "Список литературы"
            if self.REFERENCES_HEADER_PATTERN.match(text):
                references_idx = i

            # Ищем раздел "Список иллюстративного материала"
            if self.ILLUSTRATIONS_LIST_PATTERN.match(text):
                illustrations_list_idx = i

            # Ищем оглавление
            if text.lower() == "оглавление":
                in_toc = True
            elif in_toc and text:
                toc_content.append(text)
            elif in_toc and re.match(r"^(Введение|Список сокращений|Список терминов)", text):
                in_toc = False

            # Ищем ссылки на приложения в тексте
            matches = self.APPENDIX_REF_PATTERN.findall(text)
            for appendix_num in matches:
                appendix_references.append((appendix_num, i))

        # Проверяем приложения
        expected_appendix_num = 1 if appendix_number_style == "numeric" else "А"
        for para_idx, para in enumerate(document.paragraphs):
            text = para.text.strip()
            match = self.APPENDIX_HEADER_PATTERN.match(text)
            if not match:
                continue

            # Извлекаем номер приложения
            appendix_num = match.group(1)
            appendix_positions.append((appendix_num, para_idx))

            # Проверка стиля нумерации
            if appendix_number_style == "numeric" and not appendix_num.isdigit():
                errors.append(f"Приложение (параграф {para_idx+1}): Ожидается числовая нумерация (1, 2, 3), получено '{appendix_num}'")
            elif appendix_number_style == "alpha" and not re.match(r'^[А-Я]$', appendix_num):
                errors.append(f"Приложение (параграф {para_idx+1}): Ожидается буквенная нумерация (А, Б, В), получено '{appendix_num}'")

            # Проверка последовательности нумерации
            expected = str(expected_appendix_num) if appendix_number_style == "numeric" else expected_appendix_num
            if appendix_num != expected:
                errors.append(f"Приложение (параграф {para_idx+1}): Неверный номер приложения, ожидается {expected}, получено {appendix_num}")
            if appendix_number_style == "numeric":
                expected_appendix_num = int(expected_appendix_num) + 1
            else:
                expected_appendix_num = chr(ord(expected_appendix_num) + 1)

            # Проверка выравнивания заголовка "Приложение N"
            if para.alignment != WD_ALIGN_PARAGRAPH.RIGHT:
                errors.append(f"Приложение {appendix_num} (параграф {para_idx+1}): Заголовок 'Приложение {appendix_num}' должен быть выровнен по правому краю")

            # Проверка разрыва страницы перед приложением
            prev_elements = para._element.xpath('preceding-sibling::*')
            has_page_break = any(elem.tag.endswith('br') and elem.get(qn('w:type')) == 'page' for elem in prev_elements)
            if not has_page_break and para_idx > 0:
                errors.append(f"Приложение {appendix_num} (параграф {para_idx+1}): Приложение должно начинаться с новой страницы (отсутствует разрыв страницы)")

            # Проверка тематического заголовка приложения
            title_idx = para_idx + 1
            if title_idx >= len(document.paragraphs):
                errors.append(f"Приложение {appendix_num} (параграф {para_idx+1}): Отсутствует тематический заголовок после 'Приложение {appendix_num}'")
                continue
            title_para = document.paragraphs[title_idx]
            title_text = title_para.text.strip()
            if not title_text:
                errors.append(f"Приложение {appendix_num} (параграф {para_idx+1}): Отсутствует тематический заголовок после 'Приложение {appendix_num}'")
                continue

            # Проверка оформления тематического заголовка
            if not title_text[0].isupper():
                errors.append(f"Приложение {appendix_num} (параграф {title_idx+1}): Тематический заголовок '{title_text}' должен начинаться с прописной буквы")
            if title_text.endswith('.'):
                errors.append(f"Приложение {appendix_num} (параграф {title_idx+1}): Тематический заголовок '{title_text}' не должен заканчиваться точкой")
            if title_para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                errors.append(f"Приложение {appendix_num} (параграф {title_idx+1}): Тематический заголовок '{title_text}' должен быть выровнен по центру")

            # Проверка положения приложения
            if references_idx is not None and para_idx < references_idx:
                errors.append(f"Приложение {appendix_num} (параграф {para_idx+1}): Приложение должно располагаться после списка литературы")
            if illustrations_list_idx is not None and para_idx < illustrations_list_idx:
                errors.append(f"Приложение {appendix_num} (параграф {para_idx+1}): Приложение должно располагаться после списка иллюстративного материала")

        # Проверка ссылок на приложения
        referenced_appendices = set(ref[0] for ref in appendix_references)
        for appendix_num, _ in appendix_positions:
            if appendix_num not in referenced_appendices:
                errors.append(f"Приложение {appendix_num}: Отсутствует ссылка на приложение в тексте")

        # Проверка включения приложений в оглавление
        if appendix_positions:
            appendix_count = len(appendix_positions)
            if appendix_count == 1:
                # Если приложение одно, в оглавлении должно быть просто "Приложение"
                expected_toc_entry = "Приложение"
                if not any(expected_toc_entry.upper() in toc_line.upper() for toc_line in toc_content):
                    errors.append("Приложение не включено в оглавление (ожидается 'Приложение')")
            else:
                # Если приложений больше одного, проверяем их перечисление
                appendix_nums = [num for num, _ in appendix_positions]
                if appendix_count > 5:
                    # Если приложений больше 5, должен быть диапазон (например, "Приложения А–Е")
                    first_num, last_num = appendix_nums[0], appendix_nums[-1]
                    expected_toc_entry = f"Приложения {first_num}–{last_num}"
                    if not any(expected_toc_entry.upper() in toc_line.upper() for toc_line in toc_content):
                        errors.append(f"Приложения не включены в оглавление (ожидается '{expected_toc_entry}')")
                else:
                    # Если приложений 5 или меньше, они перечисляются
                    for appendix_num in appendix_nums:
                        expected_toc_entry = f"Приложение {appendix_num}"
                        if not any(expected_toc_entry.upper() in toc_line.upper() for toc_line in toc_content):
                            errors.append(f"Приложение {appendix_num} не включено в оглавление (ожидается '{expected_toc_entry}')")

        return errors